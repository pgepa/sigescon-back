"""
Gera o arquivo templates/template_fiscalizacao.docx necessário para o serviço de PDF.
Execute uma vez na raiz do projeto backend:
    python criar_template.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_font(cell, bold=False, color=None, size=10):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)


def add_section_header(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = text
    set_cell_bg(cell, "1E3A6E")
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    doc.add_paragraph()


def add_questao_sim_nao(doc, numero, texto, campo_s, campo_n, campo_d, nota=None, label_detalhe="Se não, detalhe as inconformidades"):
    table = doc.add_table(rows=3, cols=1)
    table.style = "Table Grid"

    # Header row
    header = table.cell(0, 0)
    header_text = f"{numero}. {texto}"
    if nota:
        header_text += f"\n({nota})"
    header.text = header_text
    set_cell_bg(header, "1E3A6E")
    for para in header.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
            run.bold = True

    # Radio row
    radio_cell = table.cell(1, 0)
    radio_cell.text = ""
    radio_para = radio_cell.paragraphs[0]
    radio_para.add_run(f"( {{{{ {campo_s} }}}} ) Sim    ( {{{{ {campo_n} }}}} ) Não")
    radio_para.runs[0].font.size = Pt(10)

    # Detalhe row
    detail_cell = table.cell(2, 0)
    detail_cell.text = ""
    detail_para = detail_cell.paragraphs[0]
    detail_para.add_run(f"{label_detalhe}:\n")
    detail_para.runs[0].font.size = Pt(9)
    detail_para.runs[0].italic = True
    detail_para.add_run(f"{{{{ {campo_d} }}}}")
    detail_para.runs[1].font.size = Pt(10)

    doc.add_paragraph()


def add_questao_sim_nao_na(doc, numero, texto, campo_s, campo_n, campo_a, campo_d, nota=None, aviso=None, label_detalhe="Se não, descreva a situação"):
    rows = 4 if aviso else 3
    table = doc.add_table(rows=rows, cols=1)
    table.style = "Table Grid"

    # Header
    header = table.cell(0, 0)
    header_text = f"{numero}. {texto}"
    if nota:
        header_text += f"\n({nota})"
    header.text = header_text
    set_cell_bg(header, "1E3A6E")
    for para in header.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
            run.bold = True

    row_idx = 1
    if aviso:
        aviso_cell = table.cell(row_idx, 0)
        aviso_cell.text = f"* {aviso}"
        set_cell_bg(aviso_cell, "FFFDE7")
        for para in aviso_cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                run.bold = True
        row_idx += 1

    # Radio
    radio_cell = table.cell(row_idx, 0)
    radio_cell.text = ""
    radio_para = radio_cell.paragraphs[0]
    radio_para.add_run(f"( {{{{ {campo_s} }}}} ) Sim    ( {{{{ {campo_n} }}}} ) Não    ( {{{{ {campo_a} }}}} ) Não se aplica")
    radio_para.runs[0].font.size = Pt(10)
    row_idx += 1

    # Detalhe
    detail_cell = table.cell(row_idx, 0)
    detail_cell.text = ""
    detail_para = detail_cell.paragraphs[0]
    detail_para.add_run(f"{label_detalhe}:\n")
    detail_para.runs[0].font.size = Pt(9)
    detail_para.runs[0].italic = True
    detail_para.add_run(f"{{{{ {campo_d} }}}}")
    detail_para.runs[1].font.size = Pt(10)

    doc.add_paragraph()


def main():
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Título ──────────────────────────────────────────────────────────────
    title_table = doc.add_table(rows=1, cols=1)
    title_table.style = "Table Grid"
    title_cell = title_table.cell(0, 0)
    title_cell.text = "FORMULÁRIO DE FISCALIZAÇÃO DE CONTRATO — LEI Nº 14.133/2021"
    set_cell_bg(title_cell, "1E3A6E")
    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # ── Cabeçalho do contrato ────────────────────────────────────────────────
    info_table = doc.add_table(rows=7, cols=2)
    info_table.style = "Table Grid"

    dados = [
        ("Número do Contrato:", "{{ nr_contrato }}", "Processo Administrativo (PAE):", "{{ pae }}"),
        ("Empresa Contratada:", "{{ empresa_nome }}", "CNPJ:", "{{ empresa_cnpj }}"),
        ("Objeto Contratual:", "{{ objeto }}", "", ""),
        ("Vigência (Início):", "{{ data_inicio }}", "Vigência (Término):", "{{ data_fim }}"),
        ("Valor Global:", "{{ valor_global }}", "Fiscal do Contrato:", "{{ fiscal_nome }}"),
        ("Matrícula do Fiscal:", "{{ numero_matricula }}", "", ""),
        ("Período de Fiscalização:", "De {{ periodo_de }} a {{ periodo_ate }}", "", ""),
    ]

    for i, (l1, v1, l2, v2) in enumerate(dados):
        row = info_table.rows[i]
        c0 = row.cells[0]
        c0.text = f"{l1} {v1}"
        for run in c0.paragraphs[0].runs:
            run.font.size = Pt(10)
        if l2:
            c1 = row.cells[1]
            c1.text = f"{l2} {v2}"
            for run in c1.paragraphs[0].runs:
                run.font.size = Pt(10)
        else:
            # Merge cells
            row.cells[0].merge(row.cells[1])

    doc.add_paragraph()

    # ── Seção 1: Fiscalização Técnica ────────────────────────────────────────
    add_section_header(doc, "FISCALIZAÇÃO TÉCNICA (EXECUÇÃO DO OBJETO)")

    add_questao_sim_nao(doc, 1,
        "A execução do objeto está em conformidade com o Termo de Referência, Projeto Básico ou Especificações?",
        "q1s", "q1n", "execucao_objeto_d")

    add_questao_sim_nao(doc, 2,
        "O prazo de execução/entrega está sendo cumprido?",
        "q2s", "q2n", "prazo_execucao_d")

    add_questao_sim_nao(doc, 3,
        "O nível de qualidade dos serviços/produtos atende aos padrões exigidos?",
        "q3s", "q3n", "nivel_qualidade_d")

    add_questao_sim_nao(doc, 4,
        "As medições de serviços ou entregas de produtos correspondem ao que foi realizado/entregue?",
        "q4s", "q4n", "medicoes_servicos_d")

    add_questao_sim_nao(doc, 5,
        "Houve ocorrências ou necessidade de correções na execução?",
        "q5s", "q5n", "ocorrencias_d")

    # ── Seção 2: Fiscalização Documental ────────────────────────────────────
    add_section_header(doc, "FISCALIZAÇÃO DOCUMENTAL E DE REGULARIDADE")

    add_questao_sim_nao(doc, 6,
        "A empresa contratada apresentou e manteve atualizadas todas as certidões e documentos de habilitação?",
        "q6s", "q6n", "documentos_habilitacao_d",
        nota="Conforme Art. 136 da Lei nº 14.133/21")

    add_questao_sim_nao(doc, 7,
        "Há indícios de subcontratação sem a prévia autorização da Administração Pública?",
        "q7s", "q7n", "subcontratacao_d",
        nota="Conforme Art. 122 da Lei nº 14.133/21")

    add_questao_sim_nao_na(doc, 8,
        "A empresa cumpriu com as obrigações trabalhistas, sociais e previdenciárias de seus empregados?",
        "q8s", "q8n", "q8a", "obrigacoes_empregados_d",
        nota="Conforme Art. 117 da Lei nº 14.133/21",
        aviso="PREENCHER SOMENTE QUANDO FOR CONTRATO DE MÃO-DE-OBRA TERCEIRIZADA",
        label_detalhe="Se não, especifique a irregularidade")

    add_questao_sim_nao_na(doc, 9,
        "As garantias contratuais (se houver) foram mantidas válidas e atualizadas?",
        "q9s", "q9n", "q9a", "garantias_contratuais_d",
        nota="Conforme Art. 96 da Lei nº 14.133/21")

    # ── Seção 3: Parecer ─────────────────────────────────────────────────────
    add_section_header(doc, "PARECER E RECOMENDAÇÕES DO FISCAL")

    add_questao_sim_nao(doc, 10,
        "O contrato está sendo executado de forma satisfatória no período fiscalizado?",
        "q10s", "q10n", "execucao_satisfatoria_d",
        label_detalhe="Se não, descreva a situação")

    # ── Data e assinatura ────────────────────────────────────────────────────
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.style = "Table Grid"
    sig_table.cell(0, 0).text = "Data: {{ data_relatorio }}"
    sig_table.cell(0, 1).text = "Local: ___________________________"
    sig_table.cell(1, 0).merge(sig_table.cell(1, 1))
    sig_cell = sig_table.cell(1, 0)
    sig_cell.text = ""
    sig_para = sig_cell.paragraphs[0]
    sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_para.add_run("\n\n\n_____________________________________________\n")
    sig_para.add_run("Assinatura do Fiscal do Contrato\n")
    sig_para.add_run("{{ fiscal_nome }} — Matrícula: {{ numero_matricula }}")

    for run in sig_para.runs:
        run.font.size = Pt(10)

    os.makedirs("templates", exist_ok=True)
    doc.save("templates/template_fiscalizacao.docx")
    print("Template criado com sucesso em: templates/template_fiscalizacao.docx")


if __name__ == "__main__":
    main()
